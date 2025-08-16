import os, json, re
from typing import List
from openai import OpenAI
from docx import Document

client = OpenAI(api_key=os.getenv("OPENAI_API_KEY", ""))

def extract_keywords(jd_text: str, top_k: int = 20) -> List[str]:
    stop = set("and or with the a an to for in of on & , . ( ) : ; /".split())
    words = re.findall(r"[A-Za-z][A-Za-z0-9+.#-]{2,}", jd_text or "")
    words = [w.lower() for w in words if w.lower() not in stop]
    freq = {}
    for w in words:
        freq[w] = freq.get(w,0)+1
    return [w for w,_ in sorted(freq.items(), key=lambda x: x[1], reverse=True)[:top_k]]

def generate_cover_letter(candidate_bio: dict, job: dict) -> str:
    prompt = f"""
    Write a concise, warm cover letter (<= 220 words) to the recruiter.

    Role: {job.get('title','')} at {job.get('company','')}
    Key needs: {', '.join(extract_keywords(job.get('jd_text',''))[:15])}
    Candidate summary: {candidate_bio.get('summary','')}
    Top skills: {', '.join(candidate_bio.get('skills',[])[:12])}
    Impact bullets: {', '.join(candidate_bio.get('bullets',[])[:4])}

    Requirements:
    - 3 short paragraphs max.
    - Specific references to the role/company where possible.
    - Confident, friendly tone. No fluff. End with a call-to-action.
    """
    resp = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[{"role":"user","content":prompt}],
        temperature=0.4,
    )
    return resp.choices[0].message.content.strip()

def tailor_resume(master_resume_json_path: str, job: dict, out_path: str):
    master = json.load(open(master_resume_json_path, encoding="utf-8"))
    jd_keywords = set(extract_keywords(job.get("jd_text","")))
    chosen = []
    for sec in master.get("experience", []):
        hits = [b for b in sec.get("bullets", []) if any(k in b.lower() for k in jd_keywords)]
        chosen.extend(hits[:3])
    # Build a simple DOCX
    doc = Document()
    doc.add_heading(master.get("name","Candidate Name"), 0)
    doc.add_paragraph(master.get("contact","email@example.com | github.com/you | city, state"))
    doc.add_heading("Summary", level=1)
    doc.add_paragraph(master.get("summary",""))
    doc.add_heading("Key Skills", level=1)
    doc.add_paragraph(", ".join(master.get("skills", [])))
    doc.add_heading("Selected Achievements (Tailored)", level=1)
    for b in chosen[:8]:
        doc.add_paragraph(b, style="List Bullet")
    # Optionally include full Experience/Education sections
    doc.add_heading("Experience (Recent)", level=1)
    for sec in master.get("experience", [])[:3]:
        doc.add_paragraph(f"{sec.get('company','')} — {sec.get('role','')} ({sec.get('dates','')})")
        for b in sec.get("bullets", [])[:3]:
            doc.add_paragraph(b, style="List Bullet")
    doc.add_heading("Education", level=1)
    for e in master.get("education", []):
        doc.add_paragraph(f"{e.get('school','')} — {e.get('degree','')} ({e.get('year','')})")
    doc.save(out_path)
    return out_path